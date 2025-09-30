from docx import Document
from docx.shared import Inches
import io
from typing import List, Dict, Any, Optional

class DocumentProcessor:
    """Handles reading and writing Word documents"""
    
    def __init__(self):
        pass
    
    def extract_text(self, doc_bytes: bytes) -> str:
        """Extract text content from a Word document"""
        try:
            doc = Document(io.BytesIO(doc_bytes))
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    def create_document(self, text_content: str) -> bytes:
        """Create a new Word document from text content"""
        try:
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = text_content.split('\n\n')
            
            for paragraph_text in paragraphs:
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    # Check if this looks like a table row (contains |)
                    if '|' in paragraph_text and paragraph_text.count('|') >= 2:
                        # Create table
                        cells = [cell.strip() for cell in paragraph_text.split('|')]
                        if len(cells) > 1:
                            table = doc.add_table(rows=1, cols=len(cells))
                            table.style = 'Table Grid'
                            row = table.rows[0]
                            for i, cell_text in enumerate(cells):
                                if i < len(row.cells):
                                    row.cells[i].text = cell_text
                    else:
                        # Regular paragraph
                        doc.add_paragraph(paragraph_text)
            
            # Save to bytes
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io.getvalue()
            
        except Exception as e:
            raise Exception(f"Failed to create document: {str(e)}")
    
    def extract_document_structure(self, doc_bytes: bytes) -> Dict[str, Any]:
        """Extract detailed structure information from document"""
        try:
            doc = Document(io.BytesIO(doc_bytes))
            
            structure = {
                'paragraphs': [],
                'tables': [],
                'headings': [],
                'word_count': 0,
                'paragraph_count': len(doc.paragraphs)
            }
            
            word_count = 0
            
            # Analyze paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    para_info = {
                        'index': i,
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'word_count': len(text.split())
                    }
                    
                    structure['paragraphs'].append(para_info)
                    word_count += para_info['word_count']
                    
                    # Check if it's a heading
                    if paragraph.style and 'Heading' in paragraph.style.name:
                        structure['headings'].append({
                            'level': paragraph.style.name,
                            'text': text,
                            'index': i
                        })
            
            # Analyze tables
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'content': []
                }
                
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_content.append(cell_text)
                        word_count += len(cell_text.split())
                    table_info['content'].append(row_content)
                
                structure['tables'].append(table_info)
            
            structure['word_count'] = word_count
            return structure
            
        except Exception as e:
            raise Exception(f"Failed to extract document structure: {str(e)}")
    
    def format_text_for_display(self, text: str, max_length: int = 500) -> str:
        """Format text for display in the UI"""
        if not text:
            return "No content"
        
        # Clean up text
        text = text.strip()
        
        # Truncate if too long
        if len(text) > max_length:
            text = text[:max_length] + "..."
        
        return text
    
    def validate_document(self, doc_bytes: bytes) -> Dict[str, Any]:
        """Validate that the document can be processed"""
        try:
            doc = Document(io.BytesIO(doc_bytes))
            
            validation = {
                'valid': True,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'has_content': False,
                'errors': []
            }
            
            # Check if document has any content
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    validation['has_content'] = True
                    break
            
            if not validation['has_content']:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                validation['has_content'] = True
                                break
                        if validation['has_content']:
                            break
                    if validation['has_content']:
                        break
            
            if not validation['has_content']:
                validation['errors'].append("Document appears to be empty")
            
            return validation
            
        except Exception as e:
            return {
                'valid': False,
                'errors': [f"Document validation failed: {str(e)}"]
            }
